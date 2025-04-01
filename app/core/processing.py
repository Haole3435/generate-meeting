import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))

from langchain.chains import LLMChain
from langgraph.graph import Graph
from openai import OpenAI
from docx import Document
from pydantic import BaseModel
from typing import Dict, Any
import uuid
from dotenv import load_dotenv
from pydub import AudioSegment
import tempfile
import logging
from pathlib import Path

from app.core.chains import (
    abstract_summary_chain,
    key_points_chain,
    action_items_chain,
    sentiment_chain
)
from app.core.models import ProcessedMeeting
from app.schemas.model import MeetingMinutesResponse, MeetingMinutes

# Configure logging
logger = logging.getLogger(__name__)

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

class MeetingProcessor:
    def __init__(self):
        self.workflow = Graph()
        self._setup_workflow()
        self.compiled_workflow = self.workflow.compile()

    def _setup_workflow(self):
        self.workflow.add_node("transcription", lambda state: state)
        self.workflow.add_node("abstract_summary", self._run_abstract_summary)
        self.workflow.add_node("key_points", self._run_key_points)
        self.workflow.add_node("action_items", self._run_action_items)
        self.workflow.add_node("sentiment", self._run_sentiment)

        self.workflow.set_entry_point("transcription")
        self.workflow.add_edge("transcription", "abstract_summary")
        self.workflow.add_edge("abstract_summary", "key_points")
        self.workflow.add_edge("key_points", "action_items")
        self.workflow.add_edge("action_items", "sentiment")
        self.workflow.set_finish_point("sentiment")

    def _run_abstract_summary(self, state: Dict[str, Any]):
        try:
            logger.info("Running abstract summary chain")
            state["text"] = state.get("transcription", "")
            result = abstract_summary_chain.invoke(state)
            if not hasattr(result, 'content'):
                raise ValueError("Invalid chain output format for abstract summary")
            state["abstract_summary"] = result.content
            logger.info("Abstract summary completed successfully")
            return state
        except Exception as e:
            logger.error(f"Error in abstract summary: {str(e)}", exc_info=True)
            raise

    def _run_key_points(self, state: Dict[str, Any]):
        try:
            logger.info("Running key points chain")
            state["text"] = state.get("transcription", "")
            if not state["text"]:
                raise ValueError("Empty input for key points extraction")
            
            result = key_points_chain.invoke(state)
            if not hasattr(result, 'content'):
                raise ValueError("Invalid chain output format for key points")
            
            state["key_points"] = result.content
            logger.info("Key points completed successfully")
            return state
        except Exception as e:
            logger.error(f"Error in key points: {str(e)}", exc_info=True)
            raise

    def _run_action_items(self, state: Dict[str, Any]):
        try:
            logger.info("Running action items chain")
            state["text"] = state.get("transcription", "")
            result = action_items_chain.invoke(state)
            if not hasattr(result, 'content'):
                raise ValueError("Invalid chain output format for action items")
            state["action_items"] = result.content
            logger.info("Action items completed successfully")
            return state
        except Exception as e:
            logger.error(f"Error in action items: {str(e)}", exc_info=True)
            raise

    def _run_sentiment(self, state: Dict[str, Any]):
        try:
            logger.info("Running sentiment chain")
            state["text"] = state.get("transcription", "")
            result = sentiment_chain.invoke(state)
            if not hasattr(result, 'content'):
                raise ValueError("Invalid chain output format for sentiment")
            state["sentiment"] = result.content
            logger.info("Sentiment analysis completed successfully")
            return state
        except Exception as e:
            logger.error(f"Error in sentiment analysis: {str(e)}", exc_info=True)
            raise

    def process(self, transcription: str) -> Dict[str, Any]:
        try:
            logger.info("Starting meeting processing")
            # Initialize state with transcription
            initial_state = {
                "transcription": transcription,
                "abstract_summary": "",
                "key_points": "",
                "action_items": "",
                "sentiment": ""
            }
            result = self.compiled_workflow.invoke(initial_state)
            
            # Ensure all required fields are present and properly formatted
            processed_result = {
                "transcription": transcription,  # Add transcription to output
                "general_summary": transcription,  # Using transcription as general summary
                "abstract_summary": result.get("abstract_summary", "No summary available"),
                "key_points": result.get("key_points", "").split("\n") if result.get("key_points") else [],
                "action_items": result.get("action_items", "").split("\n") if result.get("action_items") else [],
                "sentiment": result.get("sentiment", "Neutral")
            }
            logger.info("Meeting processing completed successfully")
            return processed_result
        except Exception as e:
            logger.error(f"Error in process method: {str(e)}", exc_info=True)
            raise

def convert_to_supported_format(file_path: str) -> str:
    """Convert file audio/video to supported format"""
    try:
        logger.info(f"Converting file to supported format: {file_path}")
        sound = AudioSegment.from_file(file_path)
        output_file = f"{tempfile.mktemp()}.wav"
        sound.export(output_file, format="wav")
        logger.info(f"File converted successfully to: {output_file}")
        return output_file
    except Exception as e:
        logger.error(f"Error converting file: {str(e)}", exc_info=True)
        raise ValueError(f"Cannot handle file: {str(e)}")

def process_meeting_input(input_data: str, input_type: str) -> MeetingMinutes:
    processor = MeetingProcessor()
    
    try:
        logger.info(f"Processing meeting input of type: {input_type}")
        if input_type in ["audio", "video"]:
            converted_file = convert_to_supported_format(input_data)
            try:
                with open(converted_file, "rb") as audio_file:
                    transcription_result = client.audio.transcriptions.create(
                        model="whisper-1",
                        file=audio_file
                    )
                    transcription = transcription_result.text
                logger.info("Transcription completed successfully")
            finally:
                if os.path.exists(converted_file):
                    os.remove(converted_file)
                    logger.info(f"Cleaned up converted file: {converted_file}")
        else:
            transcription = input_data
            
        logger.info(f"Transcription: {transcription[:100]}...")
        results = processor.process(transcription)
        logger.info(f"Processing results: {results}")
        return MeetingMinutes(**results)
    
    except Exception as e:
        logger.error(f"Processing failed: {str(e)}", exc_info=True)
        raise RuntimeError(f"Processing failed: {str(e)}") from e

def save_minutes_to_word(minutes: MeetingMinutes) -> str:
    try:
        logger.info("Saving minutes to Word document")
        doc = Document()
        doc.add_heading("Meeting Minutes", 0)
        
        # Add General Summary section
        doc.add_heading("General Summary", 1)
        doc.add_paragraph(minutes.general_summary)
        
        # Add Abstract Summary section
        doc.add_heading("Abstract Summary", 1)
        doc.add_paragraph(minutes.abstract_summary)
        
        # Add Key Points section
        doc.add_heading("Key Points", 1)
        for point in minutes.key_points:
            if point.strip():  # Only add non-empty points
                doc.add_paragraph(point.strip(), style="ListBullet")
        
        # Add Action Items section
        doc.add_heading("Action Items", 1)
        for item in minutes.action_items:
            if item.strip():  # Only add non-empty items
                doc.add_paragraph(item.strip(), style="ListBullet")
        
        # Add Sentiment Analysis section
        doc.add_heading("Sentiment Analysis", 1)
        doc.add_paragraph(minutes.sentiment)
        
        # Add Full Transcript section
        doc.add_heading("Full Transcript", 1)
        doc.add_paragraph(minutes.transcription)
        
        # Ensure outputs directory exists
        output_dir = Path("outputs")
        output_dir.mkdir(exist_ok=True)
        
        # Generate unique filename
        filename = f"meeting_minutes_{uuid.uuid4()}.docx"
        file_path = output_dir / filename
        
        # Save the document
        doc.save(str(file_path))
        logger.info(f"Minutes saved successfully to: {file_path}")
        return filename
    except Exception as e:
        logger.error(f"Error saving minutes to Word: {str(e)}", exc_info=True)
        raise
